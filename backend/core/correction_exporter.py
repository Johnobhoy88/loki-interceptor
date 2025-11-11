"""
Correction Exporter - Multi-format export support

Supported formats:
- JSON (structured data)
- XML (standardized)
- DOCX (Microsoft Word)
- HTML (web-ready)
- Markdown (documentation)

Features:
- Format conversion
- Styling and formatting
- Metadata embedding
- Diff/comparison views
"""

import json
import xml.etree.ElementTree as ET
from typing import Dict, Any, Optional
from datetime import datetime
from io import BytesIO


class CorrectionExporter:
    """
    Multi-format correction exporter

    Features:
    - Export to JSON, XML, DOCX, HTML, Markdown
    - Include original and corrected text
    - Highlight changes and corrections
    - Embed metadata
    """

    def __init__(self):
        """Initialize exporter"""
        self.export_handlers = {
            'json': self._export_json,
            'xml': self._export_xml,
            'docx': self._export_docx,
            'html': self._export_html,
            'markdown': self._export_markdown
        }

    async def export(
        self,
        result: Dict[str, Any],
        format: str
    ) -> Any:
        """
        Export correction result to specified format

        Args:
            result: Correction result
            format: Export format (json, xml, docx, html, markdown)

        Returns:
            Exported content (format-dependent)
        """
        format_lower = format.lower()

        if format_lower not in self.export_handlers:
            raise ValueError(f"Unsupported export format: {format}")

        handler = self.export_handlers[format_lower]
        return await handler(result)

    async def _export_json(self, result: Dict[str, Any]) -> Dict[str, Any]:
        """Export as JSON"""
        return {
            'format': 'json',
            'version': '1.0',
            'timestamp': datetime.utcnow().isoformat(),
            'original_text': result.get('original_text', ''),
            'corrected_text': result.get('corrected_text', ''),
            'corrections': result.get('corrections', []),
            'suggestions': result.get('suggestions', []),
            'metadata': {
                'issues_found': result.get('issues_found', 0),
                'issues_corrected': result.get('issues_corrected', 0),
                'improvement_score': result.get('improvement_score', 0.0),
                'algorithm_version': result.get('algorithm_version', 'unknown'),
                'pipeline_execution': result.get('pipeline_execution', {})
            }
        }

    async def _export_xml(self, result: Dict[str, Any]) -> bytes:
        """Export as XML"""
        root = ET.Element('CorrectionResult')
        root.set('version', '1.0')
        root.set('timestamp', datetime.utcnow().isoformat())

        # Original text
        original = ET.SubElement(root, 'OriginalText')
        original.text = result.get('original_text', '')

        # Corrected text
        corrected = ET.SubElement(root, 'CorrectedText')
        corrected.text = result.get('corrected_text', '')

        # Metadata
        metadata = ET.SubElement(root, 'Metadata')
        ET.SubElement(metadata, 'IssuesFound').text = str(result.get('issues_found', 0))
        ET.SubElement(metadata, 'IssuesCorrected').text = str(result.get('issues_corrected', 0))
        ET.SubElement(metadata, 'ImprovementScore').text = str(result.get('improvement_score', 0.0))
        ET.SubElement(metadata, 'AlgorithmVersion').text = result.get('algorithm_version', 'unknown')

        # Corrections
        corrections = ET.SubElement(root, 'Corrections')
        for idx, correction in enumerate(result.get('corrections', [])):
            corr_elem = ET.SubElement(corrections, 'Correction')
            corr_elem.set('id', str(idx))

            if isinstance(correction, dict):
                for key, value in correction.items():
                    if key != 'metadata':
                        elem = ET.SubElement(corr_elem, key.replace('_', '').title())
                        elem.text = str(value)

        # Convert to bytes
        tree = ET.ElementTree(root)
        xml_bytes = BytesIO()
        tree.write(xml_bytes, encoding='utf-8', xml_declaration=True)
        return xml_bytes.getvalue()

    async def _export_docx(self, result: Dict[str, Any]) -> bytes:
        """Export as DOCX (Microsoft Word)"""
        try:
            from docx import Document
            from docx.shared import RGBColor, Pt
            from docx.enum.text import WD_COLOR_INDEX
        except ImportError:
            # Fallback: return text-based format
            return self._export_docx_fallback(result)

        # Create document
        doc = Document()

        # Add title
        title = doc.add_heading('Document Correction Report', 0)

        # Add metadata
        doc.add_heading('Metadata', level=1)
        metadata_table = doc.add_table(rows=5, cols=2)
        metadata_table.style = 'Light Grid Accent 1'

        metadata_cells = [
            ('Issues Found', str(result.get('issues_found', 0))),
            ('Issues Corrected', str(result.get('issues_corrected', 0))),
            ('Improvement Score', f"{result.get('improvement_score', 0.0):.2%}"),
            ('Algorithm Version', result.get('algorithm_version', 'unknown')),
            ('Timestamp', datetime.utcnow().isoformat())
        ]

        for idx, (label, value) in enumerate(metadata_cells):
            metadata_table.rows[idx].cells[0].text = label
            metadata_table.rows[idx].cells[1].text = value

        # Add original text section
        doc.add_page_break()
        doc.add_heading('Original Text', level=1)
        original_para = doc.add_paragraph(result.get('original_text', ''))

        # Add corrected text section
        doc.add_page_break()
        doc.add_heading('Corrected Text', level=1)
        corrected_para = doc.add_paragraph(result.get('corrected_text', ''))

        # Add corrections section
        doc.add_page_break()
        doc.add_heading('Corrections Applied', level=1)

        if result.get('corrections'):
            for idx, correction in enumerate(result.get('corrections', [])[:50]):  # Limit to 50
                doc.add_heading(f'Correction #{idx + 1}', level=2)

                if isinstance(correction, dict):
                    for key, value in correction.items():
                        if key != 'metadata':
                            para = doc.add_paragraph()
                            para.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
                            para.add_run(str(value))
        else:
            doc.add_paragraph('No corrections applied.')

        # Save to bytes
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        return docx_bytes.getvalue()

    def _export_docx_fallback(self, result: Dict[str, Any]) -> bytes:
        """Fallback DOCX export (text-based)"""
        content = f"""DOCUMENT CORRECTION REPORT
{'=' * 50}

METADATA:
Issues Found: {result.get('issues_found', 0)}
Issues Corrected: {result.get('issues_corrected', 0)}
Improvement Score: {result.get('improvement_score', 0.0):.2%}
Algorithm Version: {result.get('algorithm_version', 'unknown')}
Timestamp: {datetime.utcnow().isoformat()}

{'=' * 50}

ORIGINAL TEXT:
{result.get('original_text', '')}

{'=' * 50}

CORRECTED TEXT:
{result.get('corrected_text', '')}

{'=' * 50}

CORRECTIONS APPLIED:
{len(result.get('corrections', []))} corrections

"""
        return content.encode('utf-8')

    async def _export_html(self, result: Dict[str, Any]) -> str:
        """Export as HTML"""
        original_text = result.get('original_text', '').replace('<', '&lt;').replace('>', '&gt;')
        corrected_text = result.get('corrected_text', '').replace('<', '&lt;').replace('>', '&gt;')
        corrections = result.get('corrections', [])

        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Document Correction Report</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, Cantarell, sans-serif;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background: #f5f5f5;
        }}
        .container {{
            background: white;
            padding: 30px;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #333;
            border-bottom: 3px solid #4CAF50;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #555;
            margin-top: 30px;
        }}
        .metadata {{
            background: #f9f9f9;
            padding: 15px;
            border-radius: 5px;
            margin: 20px 0;
        }}
        .metadata-row {{
            display: flex;
            justify-content: space-between;
            padding: 8px 0;
            border-bottom: 1px solid #eee;
        }}
        .metadata-label {{
            font-weight: bold;
            color: #666;
        }}
        .text-box {{
            background: #fafafa;
            padding: 20px;
            border-radius: 5px;
            border-left: 4px solid #4CAF50;
            margin: 15px 0;
            white-space: pre-wrap;
            font-family: 'Courier New', monospace;
            font-size: 14px;
            line-height: 1.6;
        }}
        .correction {{
            background: #fff3cd;
            padding: 15px;
            margin: 10px 0;
            border-left: 4px solid #ffc107;
            border-radius: 4px;
        }}
        .correction-title {{
            font-weight: bold;
            color: #856404;
            margin-bottom: 10px;
        }}
        .badge {{
            display: inline-block;
            padding: 3px 8px;
            border-radius: 3px;
            font-size: 12px;
            font-weight: bold;
        }}
        .badge-success {{
            background: #d4edda;
            color: #155724;
        }}
        .badge-info {{
            background: #d1ecf1;
            color: #0c5460;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>📄 Document Correction Report</h1>

        <div class="metadata">
            <div class="metadata-row">
                <span class="metadata-label">Issues Found:</span>
                <span class="badge badge-info">{result.get('issues_found', 0)}</span>
            </div>
            <div class="metadata-row">
                <span class="metadata-label">Issues Corrected:</span>
                <span class="badge badge-success">{result.get('issues_corrected', 0)}</span>
            </div>
            <div class="metadata-row">
                <span class="metadata-label">Improvement Score:</span>
                <span>{result.get('improvement_score', 0.0):.2%}</span>
            </div>
            <div class="metadata-row">
                <span class="metadata-label">Algorithm Version:</span>
                <span>{result.get('algorithm_version', 'unknown')}</span>
            </div>
            <div class="metadata-row">
                <span class="metadata-label">Timestamp:</span>
                <span>{datetime.utcnow().isoformat()}</span>
            </div>
        </div>

        <h2>📝 Original Text</h2>
        <div class="text-box">{original_text[:5000]}{' [TRUNCATED]' if len(original_text) > 5000 else ''}</div>

        <h2>✅ Corrected Text</h2>
        <div class="text-box">{corrected_text[:5000]}{' [TRUNCATED]' if len(corrected_text) > 5000 else ''}</div>

        <h2>🔧 Corrections Applied</h2>
        <p>Total corrections: <strong>{len(corrections)}</strong></p>
"""

        # Add corrections
        for idx, correction in enumerate(corrections[:20]):  # Limit to 20 for HTML
            if isinstance(correction, dict):
                html += f"""
        <div class="correction">
            <div class="correction-title">Correction #{idx + 1}</div>
"""
                for key, value in correction.items():
                    if key != 'metadata' and not isinstance(value, dict):
                        html += f"            <div><strong>{key.replace('_', ' ').title()}:</strong> {value}</div>\n"

                html += "        </div>\n"

        if len(corrections) > 20:
            html += f"        <p><em>... and {len(corrections) - 20} more corrections</em></p>\n"

        html += """
    </div>
</body>
</html>"""

        return html

    async def _export_markdown(self, result: Dict[str, Any]) -> str:
        """Export as Markdown"""
        corrections = result.get('corrections', [])

        md = f"""# Document Correction Report

## Metadata

| Metric | Value |
|--------|-------|
| Issues Found | {result.get('issues_found', 0)} |
| Issues Corrected | {result.get('issues_corrected', 0)} |
| Improvement Score | {result.get('improvement_score', 0.0):.2%} |
| Algorithm Version | {result.get('algorithm_version', 'unknown')} |
| Timestamp | {datetime.utcnow().isoformat()} |

---

## Original Text

```
{result.get('original_text', '')[:2000]}{' [TRUNCATED]' if len(result.get('original_text', '')) > 2000 else ''}
```

---

## Corrected Text

```
{result.get('corrected_text', '')[:2000]}{' [TRUNCATED]' if len(result.get('corrected_text', '')) > 2000 else ''}
```

---

## Corrections Applied

**Total:** {len(corrections)} corrections

"""

        # Add corrections
        for idx, correction in enumerate(corrections[:30]):  # Limit to 30
            if isinstance(correction, dict):
                md += f"\n### Correction #{idx + 1}\n\n"

                for key, value in correction.items():
                    if key != 'metadata' and not isinstance(value, dict):
                        md += f"- **{key.replace('_', ' ').title()}:** {value}\n"

        if len(corrections) > 30:
            md += f"\n*... and {len(corrections) - 30} more corrections*\n"

        md += "\n---\n\n*Generated by LOKI Correction System*\n"

        return md
